import os
import io
import json
import base64
from datetime import datetime
import requests
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import sys, os
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)
from config import settings
from database.connection import get_db, connect_to_mongo, close_mongo_connection
from models import (
    AudioTranscribeRequest,
    ComplaintGenerationRequest,
    ComplaintSaveRequest,
    UserProfileSchema
)

app = FastAPI(
    title="MANU AI API",
    description="Production-Ready Full-Stack AI Web Application",
    version="1.0.0",
    docs_url="/docs"
)

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_db_client():
    settings.validate_keys()
    await connect_to_mongo()

@app.on_event("shutdown")
async def shutdown_db_client():
    await close_mongo_connection()



# AI Draft Generator utilizing configured providers
def query_ai_model(prompt: str, system_instruction: str) -> dict:
    gemini_key = settings.GEMINI_API_KEY_STRUCTURING
    if gemini_key:
        api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_key}"
        headers = {"Content-Type": "application/json"}
        payload = {
            "contents": [{
                "parts": [{"text": f"System Directive: {system_instruction}\n\nUser Input: {prompt}"}]
            }],
            "generationConfig": {
                "responseMimeType": "application/json"
            }
        }
        try:
            response = requests.post(api_url, headers=headers, json=payload, timeout=15)
            if response.status_code == 200:
                res_json = response.json()
                text_content = res_json["candidates"][0]["content"]["parts"][0]["text"]
                return json.loads(text_content)
        except Exception as e:
            print(f"Primary AI fallback failed: {e}")

    # Ultimate dry-run structural backup to ensure reliability
    return {
        "title": "Official Complaint Notice",
        "subject": "Petition regarding reported grievances.",
        "body": f"This is an official document drafted concerning the raw input: {prompt}. Please review and update details accordingly."
    }

@app.post("/api/transcribe")
async def transcribe_audio(request: AudioTranscribeRequest):
    gemini_key = settings.GEMINI_API_KEY_VOICE
    if not gemini_key:
        raise HTTPException(status_code=500, detail="Voice API Key not configured")

    cleaned_b64 = request.audio_base64.split(",")[-1] if "," in request.audio_base64 else request.audio_base64
    
    api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_key}"
    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [{
            "parts": [
                {"text": f"Please transcribe the following audio carefully. The language is {request.language}."},
                {"inlineData": {"mimeType": "audio/webm", "data": cleaned_b64}}
            ]
        }]
    }
    try:
        response = requests.post(api_url, headers=headers, json=payload, timeout=30)
        if response.status_code == 200:
            res_json = response.json()
            text_content = res_json["candidates"][0]["content"]["parts"][0]["text"]
            return {"transcript": text_content.strip()}
        else:
            raise HTTPException(status_code=response.status_code, detail=f"Gemini API error: {response.text}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")

@app.post("/api/generate-complaint")
async def generate_complaint(request: ComplaintGenerationRequest):
    p = request.profile
    today = datetime.utcnow().strftime("%d/%m/%Y")
    place = p.district if p else "[Place]"

    profile_block = "[Your Name]\n[Door No. / House No.]\n[Street Name]\n[Ward No.] & [Zone]\n[Village / Town – Postal Code]\nMobile No.: [Mobile Number]"
    if p:
        profile_block = (
            f"{p.name}\n"
            f"{p.address}\n"
            f"{p.district}, {p.state} - {p.pincode}\n"
            f"Mobile No.: {p.phone}"
        )

    system_instruction = (
        f"You are an expert Indian government petition drafter writing in {request.language}.\n"
        "Your task: convert the citizen's raw complaint into a formal official petition letter body.\n"
        "Rules:\n"
        "1. Return ONLY valid JSON with exactly three keys: 'title', 'subject', 'body'.\n"
        "2. 'title' — short document title, e.g. 'Petition Regarding Road Damage'.\n"
        "3. 'subject' — one concise sentence starting with 'Petition requesting...' or 'Complaint regarding...'.\n"
        "4. 'body' — the full letter body text ONLY (no From/To/Date/Subject headers — those are added separately).\n"
        "   The body MUST follow this exact paragraph structure:\n"
        "   Paragraph 1: 'I am a resident living at the above-mentioned address.'\n"
        "   Paragraph 2: Describe the issue clearly, mentioning impact on senior citizens, women, and children.\n"
        "   Paragraph 3: 'Although this matter has already been brought to the attention of the concerned authorities verbally / in writing, no appropriate action has been taken so far.'\n"
        "   Paragraph 4: 'Therefore, I humbly request you to kindly give your personal attention to this matter, inspect the above-mentioned issue immediately, take the necessary action at the earliest, and provide a permanent solution in the interest of the public.'\n"
        "   Paragraph 5: 'I thank you in advance for your prompt action.'\n"
        "5. Remove all verbal fillers. Use formal, authoritative administrative language.\n"
        "6. Do NOT include salutation, sign-off, or attachments in the body — those are added separately."
    )

    prompt = (
        f"Citizen's raw complaint / spoken transcript:\n'{request.rawInput}'\n\n"
        f"Document template category: {request.templateName}\n"
        f"Citizen profile:\n"
        f"  Name: {p.name if p else '[Name]'}\n"
        f"  Address: {p.address if p else '[Address]'}, {p.district if p else '[District]'}, {p.state if p else '[State]'} - {p.pincode if p else '[Pincode]'}\n"
        f"  Phone: {p.phone if p else '[Phone]'}\n"
        f"  Language: {request.language}"
    )

    result = query_ai_model(prompt, system_instruction)
    return result

@app.get("/api/complaints")
async def list_complaints():
    db = get_db()
    if db is None:
        return []

    try:
        cursor = db.complaints.find().sort("createdAt", -1)
        complaints = await cursor.to_list(length=100)
        
        # Format for response
        for c in complaints:
            if "_id" in c:
                del c["_id"]
        return complaints
    except Exception as e:
        print(f"Error fetching complaints: {e}")
        return []

@app.post("/api/complaints", status_code=201)
async def save_complaint(request: ComplaintSaveRequest):
    db = get_db()
    if db is None:
        raise HTTPException(status_code=503, detail="Database is not configured or reachable")

    new_doc = {
        "id": f"complaint_{int(datetime.utcnow().timestamp())}",
        "templateId": request.templateId,
        "title": request.title,
        "subject": request.subject,
        "body": request.body,
        "profile": request.profile.dict() if request.profile else None,
        "evidenceImages": request.evidenceImages or [],
        "problemImages": request.problemImages or [],
        "createdAt": datetime.utcnow().isoformat()
    }

    try:
        await db.complaints.insert_one(new_doc)
        if "_id" in new_doc:
            del new_doc["_id"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save complaint: {str(e)}")

    return new_doc

@app.post("/api/generate-pdf")
async def generate_pdf(request: ComplaintSaveRequest):
    buffer = io.BytesIO()

    font_name = "Helvetica"
    font_path = os.path.join(os.path.dirname(__file__), "Pavanam-Regular.ttf")
    if os.path.exists(font_path):
        try:
            pdfmetrics.registerFont(TTFont("Pavanam", font_path))
            font_name = "Pavanam"
        except Exception as e:
            print(f"Font registration error: {e}")

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=60,
        bottomMargin=60
    )

    story = []
    styles = getSampleStyleSheet()
    p = request.profile
    today = datetime.utcnow().strftime("%d/%m/%Y")
    place = p.district if p else "[Place]"

    normal = ParagraphStyle('Normal2', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=18, spaceAfter=0)
    bold_label = ParagraphStyle('BoldLabel', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=18, spaceAfter=0)
    subject_style = ParagraphStyle('Subject', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=18, spaceAfter=6)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=20, alignment=4, spaceAfter=10)
    center_style = ParagraphStyle('Center', parent=styles['Normal'], fontName=font_name, fontSize=13, leading=20, alignment=1, spaceAfter=4)
    annexure_style = ParagraphStyle('Annexure', parent=styles['Normal'], fontName=font_name, fontSize=11, leading=18, spaceAfter=4)

    # ── Title ──────────────────────────────────────────────────────────────
    story.append(RLParagraph(f"<b>{request.title.upper()}</b>", center_style))
    story.append(Spacer(1, 18))

    # ── FROM block ─────────────────────────────────────────────────────────
    story.append(RLParagraph("<b>From:</b>", bold_label))
    if p:
        story.append(RLParagraph(p.name, normal))
        for line in p.address.split(","):
            line = line.strip()
            if line:
                story.append(RLParagraph(line, normal))
        story.append(RLParagraph(f"{p.district}, {p.state} \u2013 {p.pincode}", normal))
        story.append(RLParagraph(f"Mobile No.: {p.phone}", normal))
    else:
        for line in ["[Your Name]", "[Door No. / House No.]", "[Street Name]", "[Ward No.] & [Zone]", "[Village / Town \u2013 Postal Code]", "Mobile No.: [Mobile Number]"]:
            story.append(RLParagraph(line, normal))
    story.append(Spacer(1, 14))

    # ── TO block ───────────────────────────────────────────────────────────
    story.append(RLParagraph("<b>To:</b>", bold_label))
    story.append(RLParagraph("The Respected Officer / District Administration Office,", normal))
    story.append(RLParagraph(f"{p.district if p else '[Municipality / Office]'},", normal))
    story.append(RLParagraph(f"{p.state if p else '[District]'}.", normal))
    story.append(Spacer(1, 14))

    # ── Date & Place ───────────────────────────────────────────────────────
    story.append(RLParagraph(f"<b>Date:</b> {today}", normal))
    story.append(RLParagraph(f"<b>Place:</b> {place}", normal))
    story.append(Spacer(1, 14))

    # ── Subject ────────────────────────────────────────────────────────────
    story.append(RLParagraph(f"<b>Subject:</b> {request.subject}", subject_style))
    story.append(Spacer(1, 10))

    # ── Salutation ─────────────────────────────────────────────────────────
    story.append(RLParagraph("Respected Sir / Madam,", normal))
    story.append(Spacer(1, 10))

    # ── Body paragraphs ────────────────────────────────────────────────────
    for para in request.body.strip().split("\n"):
        para = para.strip()
        if para:
            story.append(RLParagraph(para, body_style))

    story.append(Spacer(1, 14))

    # ── Closing ────────────────────────────────────────────────────────────
    story.append(RLParagraph("Thank you.", normal))
    story.append(Spacer(1, 14))
    story.append(RLParagraph("Yours faithfully,", normal))
    story.append(Spacer(1, 30))
    story.append(RLParagraph("(Signature)", normal))
    story.append(RLParagraph(f"<b>{p.name if p else '[Your Name]'}</b>", normal))
    story.append(Spacer(1, 20))

    # ── Attachments section ────────────────────────────────────────────────
    has_images = (request.evidenceImages and len(request.evidenceImages) > 0) or \
                 (request.problemImages and len(request.problemImages) > 0)
    if has_images:
        story.append(RLParagraph("<b>Attachments (if any):</b>", bold_label))
        idx = 1
        if request.problemImages:
            for i in range(len(request.problemImages)):
                label = "Problem Site Image" if i == 0 else "Problem Area Image"
                story.append(RLParagraph(f"{idx}. {label}.", annexure_style))
                idx += 1
        if request.evidenceImages:
            story.append(RLParagraph(f"{idx}. Signature list of the public / supporting evidence.", annexure_style))
            idx += 1
            for i in range(len(request.evidenceImages) - 1):
                story.append(RLParagraph(f"{idx}. Related documents / evidence.", annexure_style))
                idx += 1
        story.append(Spacer(1, 20))

        # Embed problem images
        if request.problemImages:
            for i, img_b64 in enumerate(request.problemImages):
                try:
                    cleaned = img_b64.split(",")[-1] if "," in img_b64 else img_b64
                    img_stream = io.BytesIO(base64.b64decode(cleaned))
                    label = "Problem Site Image" if i == 0 else "Problem Area Image"
                    story.append(RLParagraph(f"<b>Annexure {i+1}: {label}</b>", annexure_style))
                    story.append(Spacer(1, 6))
                    story.append(RLImage(img_stream, width=400, height=280))
                    story.append(Spacer(1, 20))
                except Exception as ex:
                    print(f"Problem image embed error: {ex}")

        # Embed evidence images
        if request.evidenceImages:
            base_idx = len(request.problemImages or []) + 1
            for i, img_b64 in enumerate(request.evidenceImages):
                try:
                    cleaned = img_b64.split(",")[-1] if "," in img_b64 else img_b64
                    img_stream = io.BytesIO(base64.b64decode(cleaned))
                    story.append(RLParagraph(f"<b>Annexure {base_idx + i}: Supporting Evidence {i+1}</b>", annexure_style))
                    story.append(Spacer(1, 6))
                    story.append(RLImage(img_stream, width=400, height=280))
                    story.append(Spacer(1, 20))
                except Exception as ex:
                    print(f"Evidence image embed error: {ex}")

    doc.build(story)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=MANU_AI_Complaint.pdf"}
    )

@app.post("/api/generate-docx")
async def generate_docx(request: ComplaintSaveRequest):
    doc = DocxDocument()
    p = request.profile
    today = datetime.utcnow().strftime("%d/%m/%Y")
    place = p.district if p else "[Place]"

    # Page margins (1 inch = 914400 EMUs, but python-docx uses Pt/Inches)
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def add_line(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
        para = doc.add_paragraph()
        para.alignment = align
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        return para

    def add_body_para(text, size=11):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.space_before = Pt(0)
        return para

    # Title
    add_line(request.title.upper(), bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # From block
    add_line("From:", bold=True)
    if p:
        add_line(p.name)
        for line in p.address.split(","):
            line = line.strip()
            if line:
                add_line(line)
        add_line(f"{p.district}, {p.state} \u2013 {p.pincode}")
        add_line(f"Mobile No.: {p.phone}")
    else:
        for line in ["[Your Name]", "[Door No. / House No.]", "[Street Name]", "[Ward No.] & [Zone]", "[Village / Town \u2013 Postal Code]", "Mobile No.: [Mobile Number]"]:
            add_line(line)
    doc.add_paragraph()

    # To block
    add_line("To:", bold=True)
    add_line("The Respected Officer / District Administration Office,")
    add_line(f"{p.district if p else '[Municipality / Office]'},")
    add_line(f"{p.state if p else '[District]'}.")
    doc.add_paragraph()

    # Date & Place
    add_line(f"Date: {today}")
    add_line(f"Place: {place}")
    doc.add_paragraph()

    # Subject
    subj_para = doc.add_paragraph()
    subj_para.paragraph_format.space_after = Pt(6)
    r1 = subj_para.add_run("Subject: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Times New Roman"
    r2 = subj_para.add_run(request.subject)
    r2.font.size = Pt(11)
    r2.font.name = "Times New Roman"
    doc.add_paragraph()

    # Salutation
    add_line("Respected Sir / Madam,")
    doc.add_paragraph()

    # Body paragraphs
    for para in request.body.strip().split("\n"):
        para = para.strip()
        if para:
            add_body_para(para)

    doc.add_paragraph()

    # Closing
    add_line("Thank you.")
    doc.add_paragraph()
    add_line("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph()
    add_line("(Signature)")
    add_line(p.name if p else "[Your Name]", bold=True)
    doc.add_paragraph()

    # Attachments list
    has_images = (request.evidenceImages and len(request.evidenceImages) > 0) or \
                 (request.problemImages and len(request.problemImages) > 0)
    if has_images:
        add_line("Attachments (if any):", bold=True)
        idx = 1
        if request.problemImages:
            for i in range(len(request.problemImages)):
                label = "Problem Site Image" if i == 0 else "Problem Area Image"
                add_line(f"{idx}. {label}.")
                idx += 1
        if request.evidenceImages:
            add_line(f"{idx}. Signature list of the public / supporting evidence.")
            idx += 1
            for i in range(len(request.evidenceImages) - 1):
                add_line(f"{idx}. Related documents / evidence.")
                idx += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=MANU_AI_Complaint.docx"}
    )

@app.get("/health")
async def health():
    return {"status":"ok"}

@app.get("/api/health")
async def api_health():
    return {"status":"ok"}

@app.get("/api/config")
async def get_config():
    # Provide minimal, client-safe configuration the frontend needs
    return {
        "supabaseUrl": settings.SUPABASE_URL,
        "supabaseAnonKey": settings.supabase_key
    }

@app.get("/api/status")
async def get_status():
    # Quick service connectivity summary for frontend status indicators
    db_status = "offline"
    db = get_db()
    if db is not None:
        try:
            await db.command("ping")
            db_status = "connected"
        except Exception:
            pass

    hugging_status = "connected" if os.getenv("HUGGINGFACE_API_KEY") else "none"

    return {
        "mysqlDb": db_status,
        "huggingFace": hugging_status
    }

@app.post("/api/analyze-missing-details")
async def analyze_missing_details(payload: dict):
    raw = (payload or {}).get("rawInput", "") or ""
    missing = []
    low = raw.lower()
    # Very small heuristic: look for personal name and phone-like digits
    if "name" not in low and not any(word in low for word in ["mr ", "mrs ", "ms ", "name:"]):
        missing.append({"key": "name", "label": "Full Name"})
    if not any(ch.isdigit() for ch in raw):
        missing.append({"key": "phone", "label": "Phone Number"})

    needs_more = len(missing) > 0
    return {"needsMoreDetails": needs_more, "missingFields": missing}

if __name__ == "__main__":
    import uvicorn, os
    port = int(os.getenv("PORT", "8000"))
    uvicorn.run(app, host="0.0.0.0", port=port)
